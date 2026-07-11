from docx import Document
from docx.shared import Pt, Inches
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.utils import ImageReader
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import Paragraph, Frame, Spacer
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
from pathlib import Path

LOGO_PATH = Path('images/aimtolearnlogo.jpeg')
DOWNLOADS = Path('downloads')
DOWNLOADS.mkdir(exist_ok=True)

CONTENT = {
    'my-best-friend': {
        'title': 'My Best Friend',
        'text': '''
**Quote:** *“A true friend is the greatest of all blessings.”* — **François de La Rochefoucauld**

Friendship is one of the most beautiful gifts of life. It is a bond built on love, trust, care, and understanding. Everyone needs a true friend who stands by them in every situation. A best friend is someone who shares our happiness, supports us in difficult times, and makes our life more meaningful. School life becomes more enjoyable because of the friends we make, and the memories created with them remain in our hearts forever.

**Quote:** *“A friend is someone who knows all about you and still loves you.”* — **Elbert Hubbard**

My best friend is a kind, honest, and caring person. We have been together for many years, and our friendship grows stronger every day. We trust each other and respect each other's opinions. We never judge one another and always encourage each other to become better. Whether it is a happy moment or a difficult one, we are always there to support one another.

**Quote:** *“Friendship doubles our joy and divides our sorrow.”* — **Marcus Tullius Cicero**

We spend most of our time together at school. We sit in the same classroom, study together, complete assignments, and help each other understand difficult lessons. Whenever one of us faces a problem in studies, the other is always ready to explain it with patience. Learning becomes easier and more enjoyable when we work together because we motivate each other to perform our best.

**Quote:** *“The happiest moments in life are those shared with friends.”* — **Unknown**

Apart from studying, we also enjoy playing together during recess and after school. We participate in different games, sports, and school activities. Sometimes we compete with each other, but we always appreciate each other's efforts. We laugh, tell jokes, and create wonderful memories that make our friendship even stronger. These joyful moments bring happiness to our everyday lives.

**Quote:** *“The future belongs to those who believe in the beauty of their dreams.”* — **Eleanor Roosevelt**

One of the best parts of our friendship is that we dream together. We often talk about our future careers, our goals, and the kind of people we want to become. We encourage each other to work hard, stay focused, and never give up. Whenever one of us feels discouraged, the other offers hope and motivation. We believe that with dedication and mutual support, our dreams can become reality.

**Quote:** *“True friendship is built on trust, respect, and understanding.”* — **Unknown**

Like every friendship, we sometimes have small disagreements. However, we never allow misunderstandings to damage our relationship. We solve our problems through honest conversation and forgive each other quickly. This teaches us the importance of patience, kindness, and respect. Our friendship has helped us become more responsible, caring, and mature individuals.

**Quote:** *“Friends are the family we choose for ourselves.”* — **Edna Buchanan**

My best friend has taught me many valuable lessons about life. From sharing and caring to honesty and teamwork, every experience with my friend has made me a better person. We celebrate each other's achievements, support each other during failures, and stand together through every challenge. I feel truly fortunate to have such a wonderful companion who fills my life with happiness and confidence.

**Quote:** *“A sweet friendship refreshes the soul.”* — **Proverbs 27:9**

In conclusion, a best friend is one of life's greatest treasures. True friendship is not measured by the number of years spent together but by the love, trust, and unforgettable memories shared. I hope our friendship remains strong throughout our lives. No matter where the future takes us, I will always cherish the moments we spent studying together, playing together, dreaming together, and growing together. A true best friend is a priceless gift that stays in our heart forever.
'''
    },
    'television': {
        'title': 'Television',
        'text': '''
**Quote:** *“Knowledge is power.”* — **Francis Bacon**

Television is one of the most useful inventions of modern science. It is found in almost every home and is loved by people of all ages. It is not only a source of entertainment but also a great source of knowledge and information. Television brings the whole world into our living room. We can watch news, sports, movies, cartoons, documentaries, and educational programs with just the press of a button.

**Quote:** *“Learning never exhausts the mind.”* — **Leonardo da Vinci**

The first successful television was demonstrated by **John Logie Baird** in **1926**. In the beginning, televisions had only black-and-white screens. As technology improved, color televisions were introduced, followed by LED, Smart TVs, and Ultra HD televisions. Today, modern televisions provide excellent picture quality and many useful features, making learning and entertainment more enjoyable than ever.

**Quote:** *“Education is the key to success.”* — **Common Saying**

Television plays an important role in education. Students can learn many new things by watching educational channels and documentaries. Science experiments, history programs, language lessons, and nature shows make learning interesting and easy. Educational programs explain difficult concepts in a simple and engaging way. They also improve our general knowledge by teaching us about different countries, cultures, inventions, and famous personalities. Many students watch educational channels to strengthen their classroom learning and prepare for exams. In this way, television has become an important learning tool for people of all ages.

**Quote:** *“Laughter is the best medicine.”* — **Common Proverb**

One of the main reasons people watch television is entertainment. It offers different types of programs such as dramas, cartoons, movies, sports, cooking shows, quiz shows, and music programs. Every member of the family can find something they enjoy. Watching television together also gives families a chance to spend quality time with one another. Live broadcasts of cricket matches, award shows, and national celebrations make people feel connected even while sitting at home. Television provides relaxation after a busy day and helps reduce stress through enjoyable and meaningful programs.

**Quote:** *“Stay informed, stay prepared.”* — **Common Saying**

Television also keeps us updated with the latest news from around the world. We can learn about weather forecasts, scientific discoveries, sports events, and important national and international news. It also spreads awareness about health, road safety, environmental protection, and other social issues. In times of emergencies, television provides quick and reliable information to the public. Many government announcements and educational campaigns are also broadcast on television to guide and inform citizens. By watching authentic news channels, people become more aware of the world and can make better decisions in their daily lives.

**Quote:** *“Excess of everything is bad.”* — **Common Proverb**

Although television has many advantages, watching it for long hours can be harmful. It may affect our eyesight, reduce physical activity, and waste valuable time. Children may also lose interest in studies or outdoor games if they watch too much television. Therefore, we should watch television for a limited time and choose useful and educational programs. Parents should also monitor the programs their children watch to ensure they are suitable for their age. A balanced routine of study, exercise, and limited screen time helps maintain both physical and mental health.

**Quote:** *“Technology is a useful servant but a dangerous master.”* — **Christian Lous Lange**

Today, Smart TVs have made television even more useful. People can watch online classes, educational videos, documentaries, and live events through the internet. They can also connect with different learning platforms and enjoy high-quality programs. If used wisely, television can help us improve our knowledge and skills. Modern televisions also allow families to enjoy movies and educational content together, making learning more interactive and enjoyable. With continuous technological advancements, television is becoming smarter and more beneficial every year.

**Quote:** *“Use things, don't let things use you.”* — **Common Saying**

In conclusion, television is a wonderful invention that makes our lives easier, more enjoyable, and more informative. It entertains us, educates us, and keeps us connected with the world. Like every invention, it should be used wisely and in moderation. If we use television properly, it can become one of our best teachers and a valuable companion throughout life. It teaches us new ideas, broadens our thinking, and helps us understand different cultures and societies. Therefore, we should make television a source of learning and inspiration rather than merely a means of passing time.
'''
    }
}

HEADER_STYLE = ParagraphStyle(
    'Header',
    fontName='Helvetica-Bold',
    fontSize=16,
    leading=20,
    spaceAfter=12,
    alignment=TA_CENTER
)
BODY_STYLE = ParagraphStyle(
    'Body',
    fontName='Helvetica',
    fontSize=12,
    leading=18,
    alignment=TA_JUSTIFY,
    spaceAfter=12
)
QUOTE_STYLE = ParagraphStyle(
    'Quote',
    fontName='Helvetica-Oblique',
    fontSize=12,
    leading=18,
    alignment=TA_JUSTIFY,
    spaceAfter=18,
    leftIndent=18,
    rightIndent=18
)


def write_docx(filename, title, text):
    doc = Document()
    sections = doc.sections
n    section = sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = 'AIM TO LEARN'
    header_paragraph.style = doc.styles['Normal']

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = '© AIM TO LEARN'
    footer_paragraph.style = doc.styles['Normal']
    footer_paragraph.alignment = 1

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = 1
    run = title_paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph()
    for paragraph in text.strip().split('\n\n'):
        if paragraph.startswith('**Quote:**'):
            run_text = paragraph.replace('**Quote:**', '').strip()
            quote_para = doc.add_paragraph()
            quote_para.style = doc.styles['Quote'] if 'Quote' in doc.styles else doc.styles['Normal']
            quote_para.alignment = 1
            quote_run = quote_para.add_run(run_text)
            quote_run.italic = True
        else:
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            p.alignment = 3
            run = p.add_run(paragraph)
            run.font.size = Pt(12)

    doc.save(DOWNLOADS / filename)


def write_pdf(filename, title, text):
    pdf_path = DOWNLOADS / filename
    c = canvas.Canvas(str(pdf_path), pagesize=A4)
    width, height = A4
    margin = 25 * mm
    watermark_size = width * 0.5
    if LOGO_PATH.exists():
        try:
            logo = ImageReader(str(LOGO_PATH))
            c.saveState()
            c.translate((width - watermark_size) / 2, (height - watermark_size) / 2)
            c.setFillAlpha(0.1)
            c.drawImage(logo, 0, 0, watermark_size, watermark_size, mask='auto')
            c.restoreState()
        except Exception:
            pass

    c.setFont('Helvetica-Bold', 18)
    c.drawCentredString(width / 2, height - margin, title)
    c.setFont('Helvetica', 11)

    frame = Frame(margin, margin, width - 2 * margin, height - 2 * margin - 40, showBoundary=0)
    story = []
    for paragraph in text.strip().split('\n\n'):
        if paragraph.startswith('**Quote:**'):
            content = paragraph.replace('**Quote:**', '').strip()
            story.append(Spacer(1, 12))
            story.append(Paragraph(content, ParagraphStyle('quote', parent=BODY_STYLE, fontName='Helvetica-Oblique', alignment=TA_CENTER, spaceAfter=14)))
            story.append(Spacer(1, 6))
        else:
            story.append(Paragraph(paragraph, BODY_STYLE))

    frame.addFromList(story, c)
    c.showPage()
    c.save()


if __name__ == '__main__':
    for key, item in CONTENT.items():
        write_docx(f'{key}.docx', item['title'], item['text'])
        write_pdf(f'{key}.pdf', item['title'], item['text'])
